#!/usr/bin/env python3
"""Convert lab .md files to .docx with formatted title page.

Usage:
    python converter-for-clunkers.py                     # convert all files from LABS + RGR dicts
    python converter-for-clunkers.py -f file.md          # convert single file (uses LABS/RGR for metadata)
    python converter-for-clunkers.py -l list.txt         # convert files listed in list.txt

list.txt format (tab-separated):
    filename.md\tТип роботи\tНомер\tТема
    ЛБ1_Ініціація_проекту.md\tЛабораторна робота\t1\tІніціація проекту

If metadata not found (no entry in LABS/RGR, no list.txt entry), converts without title page.
"""

import argparse
import json
import shutil
import subprocess
import sys
import tempfile
import os
import re

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = os.path.join(os.getcwd(), "output")


def _find_pandoc(cfg):
    """Find pandoc: config override > PATH lookup."""
    path = cfg.get('pandoc_path') if cfg else None
    if path and os.path.isfile(path):
        return path
    found = shutil.which('pandoc')
    if found:
        return found
    raise FileNotFoundError(
        "pandoc not found. Install pandoc and add to PATH, "
        "or set \"pandoc_path\" in converter-for-clunkers.json"
    )


def make_run(p, text, size=14, bold=False, italic=False, underline=False):
    """Create a run with explicit formatting (no inheritance)."""
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    run.bold = bold
    run.italic = italic
    run.underline = underline
    return run


def add_centered(doc, text, size=14, bold=False, space_after=0, space_before=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.first_line_indent = Cm(0)
    pf.line_spacing = 1.0
    make_run(p, text, size=size, bold=bold)
    return p


def add_right(doc, text, size=14, bold=False, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    make_run(p, text, size=size, bold=bold)
    return p


def add_empty(doc, count=1, size=14):
    for _ in range(count):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        pf.first_line_indent = Cm(0)
        pf.line_spacing = 1.0
        make_run(p, '', size=size)


def add_indented(doc, indent_cm=7.0, space_after=0):
    """Paragraph with left indent (simulates right-side positioning like ref.odt)."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(indent_cm)
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    pf.first_line_indent = Cm(0)
    pf.line_spacing = 1.0
    return p


def _uline(p, text, size=12):
    """Underlined run (for handwriting blanks and filled-in text)."""
    return make_run(p, text, size=size, underline=True)


def _small_centered(doc, text, indent_cm=7.0):
    """Small centered caption like (ПІБ студента) — 8pt, indented."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.left_indent = Cm(indent_cm)
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.first_line_indent = Cm(0)
    pf.line_spacing = 1.0
    make_run(p, text, size=8)
    return p


DEFAULTS = {
    # --- Title page fields (placeholders if not set) ---
    'discipline': '[дисципліна]',
    'teacher_title': '[посада викладача]',
    'teacher_name': '[ПІБ викладача]',
    'student_name': '[ПІБ студента]',
    'student_group': '[група]',
    'student_year': '[курс]',
    'specialty': '[код та назва спеціальності]',
    'city': '[місто]',
    'year': '[рік]',
    # --- Flags ---
    'no_title': False,
    'no_toc': False,
    'keep_hr': False,
    'page_numbers': True,
    'page_number_start': 2,
    # --- Fonts ---
    'font_name': 'Times New Roman',
    'font_size': 14,
    'code_font_name': 'Courier New',
    'code_font_size': 14,
    # --- Body formatting ---
    'line_spacing': 1.08,
    'first_line_indent': 1.0,
    'heading_align': 'center',
    'table_align': 'left',
    'heading_space_before': 6,
    'heading_space_after': 0,
    # --- List formatting (cm) ---
    'list_left_indent': 1.0,
    'list_hanging_indent': 0,
    # --- Title page margins (cm) ---
    'title_margin_top': 2.0,
    'title_margin_bottom': 2.0,
    'title_margin_left': 2.0,
    'title_margin_right': 2.0,
    # --- Body margins (cm) ---
    'body_margin_top': 2.0,
    'body_margin_bottom': 2.0,
    'body_margin_left': 3.0,
    'body_margin_right': 1.5,
}

CONFIG_FILE = 'converter-for-clunkers.json'


def load_config(config_path=None):
    """Load config from JSON file, merging with defaults."""
    cfg = dict(DEFAULTS)
    path = config_path or os.path.join(os.getcwd(), CONFIG_FILE)
    if os.path.exists(path):
        with open(path, 'r', encoding='utf-8') as f:
            cfg.update(json.load(f))
    return cfg


def parse_teacher(teacher_str):
    """Split 'к.т.н., доцент каф. 603 Мандрікова Л. В.' into (title, name)."""
    # Last 3 words = name (Прізвище І. Б.)
    parts = teacher_str.rsplit(maxsplit=2)
    if len(parts) >= 3:
        name = ' '.join(parts[-2:])
        title = ' '.join(parts[:-2])
    else:
        title = ''
        name = teacher_str
    return title, name


def build_title_page(doc, lab_type, lab_num, topic, cfg):
    """Build title page matching ref.odt structure exactly."""
    INDENT = 7.0  # cm, matches ref.odt fo:text-indent="6.981cm"

    # Header block (P15 — centered, 14pt)
    add_centered(doc, 'МІНІСТЕРСТВО ОСВІТИ І НАУКИ УКРАЇНИ', 14)
    add_empty(doc, 1)
    add_centered(doc, 'Національний аерокосмічний університет', 14)
    add_centered(doc, '«Харківський авіаційний інститут»', 14)
    add_empty(doc, 1)
    add_centered(doc, 'Факультет програмної інженерії та бізнесу', 14)
    add_empty(doc, 1)
    add_centered(doc, 'Кафедра інженерії програмного забезпечення', 14)

    add_empty(doc, 5)

    # Lab title (P19 — centered, 22pt bold)
    num_display = f" №{lab_num}" if lab_num else ""
    add_centered(doc, f'{lab_type}{num_display}', 22, bold=True)

    # Discipline line (P10 — centered, 14pt normal + underlined discipline name)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    make_run(p, 'з дисципліни «', 14)
    _uline(p, f' {cfg["discipline"]} ', 14)
    make_run(p, '»', 14)

    # Subscript caption under discipline
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    run = make_run(p, '(назва дисципліни)', 14)
    run.font.subscript = True

    add_empty(doc, 6)

    # "Виконав:" line
    p = add_indented(doc, INDENT)
    make_run(p, f'Виконав: студент {cfg["student_year"]} курсу групи № ', 14)
    _uline(p, f' {cfg["student_group"]}', 14)
    _uline(p, ' ')

    # "напряму підготовки (спеціальності):" line
    p = add_indented(doc, INDENT)
    make_run(p, f'напряму підготовки (спеціальності)', 14)
    _uline(p, ' ')

    # Specialty line
    p = add_indented(doc, INDENT)
    _uline(p, cfg['specialty'])
    _uline(p, ' ')

    # (код спеціальності) — P20, 8pt, indented+centered
    _small_centered(doc, '(код спеціальності)', INDENT)

    # Student name
    p = add_indented(doc, INDENT)
    _uline(p, '\t' * 3)
    _uline(p, cfg['student_name'], 14)
    _uline(p, '\t' * 2)

    # (ПІБ студента) — P20, 8pt
    _small_centered(doc, '(ПІБ студента)', INDENT)

    add_empty(doc, 1, size=14)

    # "Прийняв:" line
    p = add_indented(doc, INDENT)
    make_run(p, 'Прийняв:', 14)
    _uline(p, f' {cfg["teacher_title"]}', 14)
    _uline(p, ' ')

    # Signature line
    p = add_indented(doc, INDENT)
    _uline(p, '\t' * 3)
    _uline(p, f' {cfg["teacher_name"]}', 14)
    _uline(p, ' ')

    # (ПІБ викладача) — P20, 8pt
    _small_centered(doc, '(ПІБ викладача)', INDENT)

    add_empty(doc, 1, size=14)

    # Grade lines
    p = add_indented(doc, INDENT)
    make_run(p, 'Національна шкала: __________', 14)
    _uline(p, ' ')

    p = add_indented(doc, INDENT)
    make_run(p, 'Кількість балів:\t', 14)
    _uline(p, ' ')

    p = add_indented(doc, INDENT)
    make_run(p, 'Оцінка:  ECTS\t', 14)
    _uline(p, ' ')
    
    add_empty(doc, 6)

    # City/year + page break in same paragraph
    p = add_centered(doc, f'{cfg["city"]} – {cfg["year"]}', 14)
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)


def build_toc(doc):
    """Insert auto-generated Table of Contents field."""
    add_centered(doc, 'ЗМІСТ', 14, bold=True)

    # Insert TOC field — Word will update it on open
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)

    # Begin complex field: TOC \o "1-3" \h \z \u
    r1 = p.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    r1._element.append(fld_begin)

    r2 = p.add_run()
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = r' TOC \o "1-3" \h \z \u '
    r2._element.append(instr)

    r3 = p.add_run()
    fld_separate = OxmlElement('w:fldChar')
    fld_separate.set(qn('w:fldCharType'), 'separate')
    r3._element.append(fld_separate)

    r4 = p.add_run('(Оновіть зміст: ПКМ → Оновити поле)')
    r4.font.size = Pt(11)
    r4.font.color.rgb = None

    r5 = p.add_run()
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r5._element.append(fld_end)

    # Page break after TOC
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)


def strip_first_headings(content):
    """Remove first # line (document title, already on title page).

    Only the first H1 is stripped. The first ## is kept — it may be a
    numbered section heading (e.g. '## 1 Introduction') that carries
    its own content and tables.
    """
    lines = content.split('\n')
    result = []
    skipped_h1 = False
    for line in lines:
        if not skipped_h1 and line.startswith('# ') and not line.startswith('## '):
            skipped_h1 = True
            continue
        if skipped_h1 and not result and line.strip() == '':
            continue
        result.append(line)
    return '\n'.join(result)


def shift_headings(content):
    """Shift heading levels so the smallest heading becomes # (Heading 1).

    Auto-detects the minimum heading level in the content and shifts all
    headings down so that minimum becomes level 1.  For example:
      - If smallest is ### (3): shift by 2 → ### -> #, #### -> ##
      - If smallest is ## (2):  shift by 1 → ## -> #, ### -> ##
      - If smallest is # (1):   no shift needed
    """
    lines = content.split('\n')
    # Find minimum heading level present
    min_level = 7
    for line in lines:
        match = re.match(r'^(#{1,6})\s', line)
        if match:
            min_level = min(min_level, len(match.group(1)))
    if min_level >= 7:
        return content  # no headings found
    shift = min_level - 1
    if shift == 0:
        return content
    result = []
    for line in lines:
        if line.startswith('#'):
            match = re.match(r'^(#{1,6})\s', line)
            if match:
                level = len(match.group(1))
                new_level = max(1, level - shift)
                line = '#' * new_level + line[level:]
        result.append(line)
    return '\n'.join(result)


def _is_code_block(para):
    """Detect code block paragraphs from pandoc output.

    After element copy, python-docx can't resolve 'Source Code' style by name,
    but raw XML pStyle value 'SourceCode' survives. Also checks for pandoc
    syntax-highlighting rStyle tokens on runs (*Tok, VerbatimChar).
    """
    # Check 1: raw XML pStyle = 'SourceCode'
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None:
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is not None and pStyle.get(qn('w:val')) == 'SourceCode':
            return True

    # Check 2: runs have pandoc syntax-highlighting rStyle values
    PANDOC_CODE_RSTYLES = {
        'NormalTok', 'KeywordTok', 'DataTypeTok', 'DecValTok', 'BaseNTok',
        'FloatTok', 'CharTok', 'StringTok', 'CommentTok', 'OtherTok',
        'AlertTok', 'FunctionTok', 'RegionMarkerTok', 'ErrorTok',
        'OperatorTok', 'BuiltInTok', 'ExtensionTok', 'PreprocessorTok',
        'AttributeTok', 'DocumentationTok', 'AnnotationTok',
        'CommentVarTok', 'VariableTok', 'ControlFlowTok', 'ConstantTok',
        'SpecialCharTok', 'SpecialStringTok', 'ImportTok', 'InformationTok',
        'WarningTok', 'VerbatimChar',
    }
    runs = para.runs
    if not runs:
        return False
    code_runs = 0
    for run in runs:
        rPr = run._element.find(qn('w:rPr'))
        if rPr is not None:
            rStyle = rPr.find(qn('w:rStyle'))
            if rStyle is not None and rStyle.get(qn('w:val')) in PANDOC_CODE_RSTYLES:
                code_runs += 1
    # If majority of runs have code rStyles, it's a code block
    return code_runs > 0 and code_runs >= len(runs) // 2


def convert(filename, lab_type, lab_num, topic, src_dir, cfg=None):
    cfg = cfg or load_config()
    src_path = os.path.join(src_dir, filename)
    if not os.path.exists(src_path):
        print(f"  SKIP: {src_path} not found")
        return

    with open(src_path, "r", encoding="utf-8") as f:
        content = f.read()

    content = strip_first_headings(content)
    content = shift_headings(content)
    # Remove horizontal rules (--- lines) unless keep_hr is set
    if not cfg.get('keep_hr', False):
        content = re.sub(r'^\s*-{3,}\s*$', '', content, flags=re.MULTILINE)
    # Auto-number tables and figures: section.element format (e.g. Таблиця 1.2)
    # Tracks current Heading 1 section number.
    # Extracts explicit number from heading text (e.g. "# 2 Planning" -> 2),
    # falls back to auto-increment when heading has no leading number.
    section_num = 0
    table_in_section = 0
    figure_in_section = 0
    lines = content.split('\n')
    result_lines = []
    for line in lines:
        h1_match = re.match(r'^#\s+(.*)', line)
        if h1_match:
            num_match = re.match(r'^(\d+)', h1_match.group(1).strip())
            if num_match:
                section_num = int(num_match.group(1))
            else:
                section_num += 1
            table_in_section = 0
            figure_in_section = 0
        m_table = re.match(r'^Таблиця:\s*(.+)$', line)
        if m_table:
            table_in_section += 1
            desc = m_table.group(1).strip()
            line = f'Таблиця {section_num}.{table_in_section} — {desc}'
        m_fig = re.match(r'^Рисунок:\s*(.+)$', line)
        if m_fig:
            figure_in_section += 1
            desc = m_fig.group(1).strip()
            line = f'Рисунок {section_num}.{figure_in_section} — {desc}'
        result_lines.append(line)
    content = '\n'.join(result_lines)
    # Strip alt text from images — captions handled by "Рисунок:" lines
    content = re.sub(r'!\[[^\]]*\]\(', '![](', content)
    # Convert <!-- left --> / <!-- /left --> to marker paragraphs for post-processing
    content = re.sub(r'<!--\s*left\s*-->', 'ALIGN_LEFT_BEGIN', content)
    content = re.sub(r'<!--\s*/left\s*-->', 'ALIGN_LEFT_END', content)

    os.makedirs(OUT_DIR, exist_ok=True)
    base = os.path.splitext(os.path.basename(filename))[0]
    out_path = os.path.join(OUT_DIR, f"{base}.docx")

    # Step 1: Pandoc converts body MD -> temporary DOCX
    with tempfile.NamedTemporaryFile(mode="w", suffix=".md", encoding="utf-8", delete=False) as tmp_md:
        tmp_md.write(content)
        tmp_md_path = tmp_md.name

    tmp_docx_path = tmp_md_path.replace('.md', '.docx')

    try:
        # --resource-path so pandoc finds images relative to source file
        resource_dir = os.path.dirname(os.path.abspath(src_path))
        pandoc = _find_pandoc(cfg)
        cmd = [pandoc, tmp_md_path, "-o", tmp_docx_path,
               f"--resource-path={resource_dir}"]
        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode != 0:
            print(f"  ERROR pandoc: {result.stderr}")
            return

        # Step 2: Open pandoc docx, prepend title page
        body_doc = Document(tmp_docx_path)

        # Create new document with title page
        final_doc = Document()

        font_name = cfg['font_name']
        font_size = cfg['font_size']
        h_align = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER,
                    'right': WD_ALIGN_PARAGRAPH.RIGHT}.get(cfg['heading_align'], WD_ALIGN_PARAGRAPH.LEFT)
        no_title = cfg.get('no_title', False)

        # Set default font
        style = final_doc.styles['Normal']
        style.font.name = font_name
        style.font.size = Pt(font_size)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style.paragraph_format.first_line_indent = Cm(cfg['first_line_indent'])
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.line_spacing = cfg['line_spacing']

        # Set heading styles
        for i in range(1, 4):
            heading_style = final_doc.styles[f'Heading {i}']
            heading_style.font.name = font_name
            heading_style.font.size = Pt(font_size)
            heading_style.font.bold = True
            heading_style.font.color.rgb = None
            heading_style.paragraph_format.alignment = h_align
            heading_style.paragraph_format.first_line_indent = Cm(0)
            heading_style.paragraph_format.space_before = Pt(cfg['heading_space_before'])
            heading_style.paragraph_format.space_after = Pt(cfg['heading_space_after'])

        # Title page section margins
        section = final_doc.sections[0]
        section.top_margin = Cm(cfg['title_margin_top'])
        section.bottom_margin = Cm(cfg['title_margin_bottom'])
        section.left_margin = Cm(cfg['title_margin_left'])
        section.right_margin = Cm(cfg['title_margin_right'])

        # Build title page
        if not no_title and lab_type:
            build_title_page(final_doc, lab_type, lab_num, topic, cfg)

            if not cfg.get('no_toc', False):
                build_toc(final_doc)

        # Title page section: no page numbers
        title_section = final_doc.sections[0]
        title_section.different_first_page_header_footer = False

        # Body section margins
        new_section = final_doc.add_section()
        new_section.top_margin = Cm(cfg['body_margin_top'])
        new_section.bottom_margin = Cm(cfg['body_margin_bottom'])
        new_section.left_margin = Cm(cfg['body_margin_left'])
        new_section.right_margin = Cm(cfg['body_margin_right'])

        # Page numbering
        if cfg.get('page_numbers', True):
            page_start = cfg.get('page_number_start', 2)
            sectPr = new_section._sectPr
            pgNumType = OxmlElement('w:pgNumType')
            pgNumType.set(qn('w:start'), str(page_start))
            sectPr.append(pgNumType)

            footer = new_section.footer
            footer.is_linked_to_previous = False
            fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fp.paragraph_format.space_after = Pt(0)
            run = fp.add_run()
            fld_begin = OxmlElement('w:fldChar')
            fld_begin.set(qn('w:fldCharType'), 'begin')
            run._element.append(fld_begin)
            run2 = fp.add_run()
            instr = OxmlElement('w:instrText')
            instr.set(qn('xml:space'), 'preserve')
            instr.text = ' PAGE '
            run2._element.append(instr)
            run3 = fp.add_run()
            fld_sep = OxmlElement('w:fldChar')
            fld_sep.set(qn('w:fldCharType'), 'separate')
            run3._element.append(fld_sep)
            run4 = fp.add_run(str(page_start))
            run4.font.size = Pt(12)
            run4.font.name = font_name
            run5 = fp.add_run()
            fld_end = OxmlElement('w:fldChar')
            fld_end.set(qn('w:fldCharType'), 'end')
            run5._element.append(fld_end)

        # Step 3: Copy image parts from pandoc output to final doc
        # Images are stored as relationships in the document part.
        # We must copy them with matching rIds so XML references stay valid.
        body_part = body_doc.part
        final_part = final_doc.part
        for rel in body_part.rels.values():
            if "image" in rel.reltype:
                final_part.rels.add_relationship(rel.reltype, rel.target_part, rel.rId)

        # Step 3a: Copy numbering definitions (bullet/numbered list formats)
        # Without this, all lists render as numbered because numId references
        # in copied paragraphs can't resolve to their abstractNum definitions.
        try:
            body_numbering = body_doc.part.numbering_part.numbering_definitions._numbering
            # Ensure final_doc has a numbering part
            try:
                final_numbering = final_doc.part.numbering_part.numbering_definitions._numbering
            except Exception:
                # No numbering part yet — create one by accessing it
                from docx.opc.constants import RELATIONSHIP_TYPE as RT
                from docx.parts.numbering import NumberingPart
                from docx.opc.part import PartFactory
                numbering_part = NumberingPart.new()
                final_doc.part.relate_to(numbering_part, RT.NUMBERING)
                final_numbering = numbering_part.numbering_definitions._numbering

            # Copy all abstractNum elements, adjusting indentation
            list_left = int(cfg['list_left_indent'] * 360 / 0.635)  # cm -> twips
            list_hang = int(cfg['list_hanging_indent'] * 360 / 0.635)  # cm -> twips
            for absNum in body_numbering.findall(qn('w:abstractNum')):
                # Override indent per level to match config
                for lvl in absNum.findall(qn('w:lvl')):
                    ilvl = int(lvl.get(qn('w:ilvl'), '0'))
                    pPr = lvl.find(qn('w:pPr'))
                    if pPr is not None:
                        ind = pPr.find(qn('w:ind'))
                        if ind is not None:
                            level_left = list_left + ilvl * list_left
                            ind.set(qn('w:left'), str(level_left))
                            if list_hang == 0:
                                # Remove hanging indent — number/bullet flush with text
                                if ind.get(qn('w:hanging')) is not None:
                                    del ind.attrib[qn('w:hanging')]
                            else:
                                ind.set(qn('w:hanging'), str(list_hang))
                final_numbering.append(absNum)
            # Copy all num elements (map numId -> abstractNumId)
            for num in body_numbering.findall(qn('w:num')):
                final_numbering.append(num)
        except Exception:
            pass  # No numbering in source — nothing to copy

        # Step 3b: Convert inline images to "Top and Bottom" wrapping
        # OOXML requires strict child order in wp:anchor:
        #   simplePos, positionH, positionV, extent, effectExtent,
        #   wrapTopAndBottom, docPr, cNvGraphicFramePr, graphic
        WP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
        for drawing in body_doc.element.body.iter(qn('w:drawing')):
            for inline in drawing.findall(f'{{{WP_NS}}}inline'):
                anchor = OxmlElement('wp:anchor')
                # Required anchor attributes
                anchor.set('distT', '0')
                anchor.set('distB', '0')
                anchor.set('distL', '0')
                anchor.set('distR', '0')
                anchor.set('behindDoc', '0')
                anchor.set('locked', '0')
                anchor.set('layoutInCell', '1')
                anchor.set('allowOverlap', '1')
                anchor.set('simplePos', '0')
                anchor.set('relativeHeight', '0')
                # 1. simplePos
                simplePos = OxmlElement('wp:simplePos')
                simplePos.set('x', '0')
                simplePos.set('y', '0')
                anchor.append(simplePos)
                # 2. positionH — center relative to margin
                posH = OxmlElement('wp:positionH')
                posH.set('relativeFrom', 'margin')
                align_el = OxmlElement('wp:align')
                align_el.text = 'center'
                posH.append(align_el)
                anchor.append(posH)
                # 3. positionV
                posV = OxmlElement('wp:positionV')
                posV.set('relativeFrom', 'paragraph')
                posOffset = OxmlElement('wp:posOffset')
                posOffset.text = '0'
                posV.append(posOffset)
                anchor.append(posV)
                # 4-5. extent + effectExtent from inline
                for tag in ('extent', 'effectExtent'):
                    el = inline.find(f'{{{WP_NS}}}{tag}')
                    if el is not None:
                        anchor.append(el)
                # 6. wrapTopAndBottom
                anchor.append(OxmlElement('wp:wrapTopAndBottom'))
                # 7-9. docPr, cNvGraphicFramePr, graphic from inline
                for child in list(inline):
                    anchor.append(child)
                drawing.replace(inline, anchor)

        # Step 3c: Copy body elements from pandoc output
        for element in body_doc.element.body:
            final_doc.element.body.append(element)

        # Step 4: Post-process body — fix fonts, captions, alignment, remove horizontal rules
        left_align = False
        markers_to_remove = []
        for para in final_doc.paragraphs:
            # Fix heading fonts: pandoc uses Calibri
            if para.style and para.style.name and para.style.name.startswith('Heading'):
                para.alignment = h_align
                for run in para.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
                    run.bold = True
                    run.font.color.rgb = None
                # Page break before "Джерела" heading
                if para.text.strip().lower() in ('джерела', 'список джерел', 'література'):
                    para.paragraph_format.page_break_before = True
            # Code blocks: left-align, no indent, monospace
            elif _is_code_block(para):
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.first_line_indent = Cm(0)
                for run in para.runs:
                    run.font.name = cfg['code_font_name']
                    run.font.size = Pt(cfg['code_font_size'])
            # Fix body text font
            elif para.style and para.style.name in ('Normal', 'Body Text', 'First Paragraph'):
                for run in para.runs:
                    if run.font.name and run.font.name != font_name:
                        run.font.name = font_name

            # List items: indentation is handled at numbering definition level
            # (Step 3a), not at paragraph level — paragraph w:ind is ignored
            # when w:numPr is present and the numbering level defines its own indent.

            # Table/figure captions — no indent, keep with next
            # Tables: left-aligned, Figures: centered
            text = para.text.strip()
            if re.match(r'^Таблиця\s+\d+\.\d+\s+[—–-]', text):
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.first_line_indent = Cm(0)
                para.paragraph_format.keep_with_next = True
            elif re.match(r'^Рисунок\s+\d+\.\d+\s+[—–-]', text):
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.first_line_indent = Cm(0)
                para.paragraph_format.keep_with_next = True

            # Handle <!-- left --> / <!-- /left --> alignment markers
            if para.text.strip() == 'ALIGN_LEFT_BEGIN':
                left_align = True
                markers_to_remove.append(para)
            elif para.text.strip() == 'ALIGN_LEFT_END':
                left_align = False
                markers_to_remove.append(para)
            elif left_align:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Remove horizontal rules (pBdr elements from markdown ---)
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                pBdr = pPr.find(qn('w:pBdr'))
                if pBdr is not None:
                    pPr.remove(pBdr)

        # Remove alignment marker paragraphs
        for para in markers_to_remove:
            para._element.getparent().remove(para._element)

        # Step 5: Fix tables — borders, alignment, remove inherited first-line indent
        tbl_align = cfg.get('table_align', 'left').lower()
        for table in final_doc.tables:
            # Set thin black borders on entire table
            tbl = table._tbl
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            # Table alignment: left/center/full (full = left + 100% width)
            jc = tblPr.find(qn('w:jc'))
            if jc is None:
                jc = OxmlElement('w:jc')
                tblPr.append(jc)
            jc.set(qn('w:val'), 'left' if tbl_align == 'full' else tbl_align)
            tblBorders = OxmlElement('w:tblBorders')
            for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                border = OxmlElement(f'w:{edge}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')  # 0.5pt
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tblBorders.append(border)
            # Remove existing borders if any
            existing = tblPr.find(qn('w:tblBorders'))
            if existing is not None:
                tblPr.remove(existing)
            tblPr.append(tblBorders)

            # Table width: full = force 100%, otherwise keep pandoc's width
            if tbl_align == 'full':
                tblW = tblPr.find(qn('w:tblW'))
                if tblW is None:
                    tblW = OxmlElement('w:tblW')
                    tblPr.append(tblW)
                tblW.set(qn('w:type'), 'pct')
                tblW.set(qn('w:w'), '5000')  # 5000 = 100% in pct units

            # Set layout to fixed (allows manual column resize in Word)
            tblLayout = tblPr.find(qn('w:tblLayout'))
            if tblLayout is None:
                tblLayout = OxmlElement('w:tblLayout')
                tblPr.append(tblLayout)
            tblLayout.set(qn('w:type'), 'fixed')

            # Prevent row split across pages, fix cell formatting
            for row in table.rows:
                trPr = row._tr.find(qn('w:trPr'))
                if trPr is None:
                    trPr = OxmlElement('w:trPr')
                    row._tr.insert(0, trPr)
                cantSplit = OxmlElement('w:cantSplit')
                trPr.append(cantSplit)

                for cell in row.cells:
                    for para in cell.paragraphs:
                        para.paragraph_format.first_line_indent = Cm(0)
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        try:
            final_doc.save(out_path)
        except PermissionError:
            print(f"  ERROR: Cannot save '{out_path}' — file is open. Close it and retry.")
            sys.exit(1)
        print(f"  OK: {out_path}")

    finally:
        if os.path.exists(tmp_md_path):
            os.unlink(tmp_md_path)
        if os.path.exists(tmp_docx_path):
            os.unlink(tmp_docx_path)


def load_list_file(list_path):
    """Load file list from tab-separated text file.

    Format per line: filename<TAB>type<TAB>number<TAB>topic
    Lines starting with # are skipped.
    """
    entries = {}
    with open(list_path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith('#'):
                continue
            parts = line.split('\t')
            if len(parts) >= 4:
                entries[parts[0]] = (parts[1], parts[2], parts[3])
            elif len(parts) == 1:
                # Just filename, no metadata
                entries[parts[0]] = None
    return entries


def convert_file(filename, metadata, src_dir, cfg=None):
    """Convert single file. metadata = (type, num, topic) or None."""
    if metadata:
        lab_type, num, topic = metadata
    else:
        lab_type, num, topic = '', '', ''
    convert(filename, lab_type, num, topic, src_dir, cfg=cfg)


def _parse_md_filename(filename):
    """Try to extract metadata (type, num, topic) from filename.

    Patterns:
        ЛБ1_Ініціація_проекту.md  -> ("Лабораторна робота", "1", "Ініціація проекту")
        РГР_Скрам_критика.md      -> ("Розрахунково-графічна робота", "", "Скрам критика")
        anything_else.md          -> ("", "", "")
    """
    base = os.path.splitext(filename)[0]
    # Known prefixes
    prefixes = {
        'ЛБ': 'Лабораторна робота',
        'ЛР': 'Лабораторна робота',
        'РГР': 'Розрахунково-графічна робота',
        'КР': 'Курсова робота',
        'ПР': 'Практична робота',
    }
    for prefix, lab_type in prefixes.items():
        if base.startswith(prefix):
            rest = base[len(prefix):]
            # Try to extract number: ЛБ1_Topic or ЛБ_Topic
            m = re.match(r'^(\d+)[_\s\-]*(.*)', rest)
            if m:
                num = m.group(1)
                topic = re.sub(r'[_\s\-]+', ' ', m.group(2)).strip()
            else:
                num = ''
                topic = re.sub(r'[_\s\-]+', ' ', rest).strip()
            return lab_type, num, topic
    return '', '', ''


def scan_md_files(config_path=None):
    """Scan cwd for .md files, parse metadata from filenames, write to config."""
    src_dir = os.getcwd()
    cfg_path = config_path or os.path.join(src_dir, CONFIG_FILE)

    # Load existing config or start fresh
    existing = {}
    if os.path.exists(cfg_path):
        with open(cfg_path, 'r', encoding='utf-8') as f:
            existing = json.load(f)

    # Index existing files entries by filename for merging
    existing_files = {e['file']: e for e in existing.get('files', [])}

    # Scan for .md files
    md_files = sorted(f for f in os.listdir(src_dir)
                      if f.endswith('.md') and not f.startswith('.'))

    # Also scan subdirectories one level deep
    for d in sorted(os.listdir(src_dir)):
        subdir = os.path.join(src_dir, d)
        if os.path.isdir(subdir) and not d.startswith('.') and d != 'output':
            for f in sorted(os.listdir(subdir)):
                if f.endswith('.md') and not f.startswith('.'):
                    md_files.append(os.path.join(d, f))

    new_files = []
    added = 0
    for md in md_files:
        basename = os.path.basename(md)
        # Skip config/docs files
        if basename == CONFIG_FILE.replace('.json', '.md'):
            continue
        if md in existing_files:
            # Keep existing entry (user may have edited metadata)
            new_files.append(existing_files[md])
        else:
            lab_type, num, topic = _parse_md_filename(basename)
            new_files.append({
                'file': md,
                'type': lab_type,
                'num': num,
                'topic': topic,
            })
            added += 1

    existing['files'] = new_files

    with open(cfg_path, 'w', encoding='utf-8') as f:
        json.dump(existing, f, ensure_ascii=False, indent=4)

    print(f"Scan: {len(new_files)} .md files ({added} new) -> {cfg_path}")
    for entry in new_files:
        marker = '*' if entry['file'] not in existing_files else ' '
        print(f"  {marker} {entry['file']}")


def create_config(config_path=None):
    """Create a starter config file with all defaults in the current directory."""
    cfg_path = config_path or os.path.join(os.getcwd(), CONFIG_FILE)
    if os.path.exists(cfg_path):
        print(f"Config already exists: {cfg_path}")
        print("Delete it first or use -c to specify a different path.")
        return

    cfg = dict(DEFAULTS)
    cfg['files'] = []

    with open(cfg_path, 'w', encoding='utf-8') as f:
        json.dump(cfg, ensure_ascii=False, indent=4, fp=f)

    print(f"Created: {cfg_path}")
    print("Edit it to fill in discipline, teacher, student info, then run 'scan' to discover .md files.")


def main():
    parser = argparse.ArgumentParser(
        prog='converter-for-clunkers',
        description=(
            'Markdown to DOCX converter with Ukrainian university title page.\n'
            'Converts .md lab reports to formatted .docx files with title page,\n'
            'table of contents, page numbering, and proper academic formatting.\n'
            '\n'
            'Requires: pandoc, python-docx'
        ),
        epilog=(
            'commands:\n'
            '  create-config   Generate starter config (converter-for-clunkers.json)\n'
            '                  with all defaults in the current directory\n'
            '  scan            Discover .md files in current directory and add them\n'
            '                  to the config "files" list\n'
            '\n'
            'workflow:\n'
            '  1. cd into your project directory\n'
            '  2. converter-for-clunkers.py create-config\n'
            '  3. Edit converter-for-clunkers.json (discipline, teacher, student...)\n'
            '  4. converter-for-clunkers.py scan\n'
            '  5. converter-for-clunkers.py              # convert all\n'
            '     converter-for-clunkers.py -f file.md   # convert one\n'
            '\n'
            'config (converter-for-clunkers.json):\n'
            '  Title page:  discipline, teacher_title, teacher_name,\n'
            '               student_name, student_group, student_year,\n'
            '               specialty, city, year\n'
            '  Flags:       no_title, no_toc, keep_hr, page_numbers\n'
            '  Fonts:       font_name, font_size, code_font_name, code_font_size\n'
            '  Layout:      line_spacing, first_line_indent, heading_align,\n'
            '               table_align, body_margin_*, title_margin_*\n'
            '  Files:       [{file, type, num, topic}, ...]\n'
        ),
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )

    parser.add_argument('command', nargs='?', default=None, metavar='COMMAND',
                        help='create-config | scan (see below)')

    convert_group = parser.add_argument_group('conversion target (pick one)')
    convert_group.add_argument('-f', '--file', metavar='MD',
                               help='convert single .md file')
    convert_group.add_argument('-L', '--list', metavar='TXT',
                               help='tab-separated list: file<TAB>type<TAB>num<TAB>topic')

    toggle_group = parser.add_argument_group('toggle flags (override config)')
    toggle_group.add_argument('-n', '--no-title', action='store_true', default=None,
                              help='skip title page')
    toggle_group.add_argument('--with-title', action='store_true', default=None,
                              help='force title page')
    toggle_group.add_argument('-H', '--keep-hr', action='store_true', default=None,
                              help='keep horizontal rules (--- lines)')
    toggle_group.add_argument('--no-hr', action='store_true', default=None,
                              help='remove horizontal rules')
    toggle_group.add_argument('--no-toc', action='store_true', default=None,
                              help='skip table of contents')
    toggle_group.add_argument('--with-toc', action='store_true', default=None,
                              help='force table of contents')
    toggle_group.add_argument('--no-page-numbers', action='store_true', default=None,
                              help='disable page numbering')
    toggle_group.add_argument('--with-page-numbers', action='store_true', default=None,
                              help='force page numbering')
    toggle_group.add_argument('--table-align', choices=['left', 'center', 'full'], default=None,
                              help='table alignment (default: left)')

    meta_group = parser.add_argument_group('metadata overrides (override config per-run)')
    meta_group.add_argument('-c', '--config', metavar='JSON',
                            help='path to config file (default: converter-for-clunkers.json next to script)')
    meta_group.add_argument('-d', '--discipline', metavar='NAME',
                            help='discipline name')
    meta_group.add_argument('-p', '--teacher', metavar='"TITLE NAME"',
                            help='teacher title + name')
    meta_group.add_argument('-s', '--student', metavar='NAME',
                            help='student name')
    meta_group.add_argument('-t', '--type', dest='lab_type', metavar='TYPE',
                            help='work type (e.g. "Лабораторна робота")')
    meta_group.add_argument('-N', '--num', metavar='N',
                            help='work number')
    meta_group.add_argument('-T', '--topic', metavar='TEXT',
                            help='work topic')

    args = parser.parse_args()

    # Handle subcommands
    KNOWN_COMMANDS = ('create-config', 'scan')
    if args.command is not None and args.command not in KNOWN_COMMANDS:
        parser.error(f"unknown command: '{args.command}' (choose from: {', '.join(KNOWN_COMMANDS)})")
    if args.command == 'create-config':
        create_config(args.config)
        return
    if args.command == 'scan':
        scan_md_files(args.config)
        return

    src_dir = os.getcwd()

    # Load config from file, then apply CLI overrides
    cfg = load_config(args.config)
    # Boolean flags: --no-X / --with-X pairs override config in both directions
    if args.no_title:
        cfg['no_title'] = True
    elif args.with_title:
        cfg['no_title'] = False
    if args.keep_hr:
        cfg['keep_hr'] = True
    elif args.no_hr:
        cfg['keep_hr'] = False
    if args.no_toc:
        cfg['no_toc'] = True
    elif args.with_toc:
        cfg['no_toc'] = False
    if args.no_page_numbers:
        cfg['page_numbers'] = False
    elif args.with_page_numbers:
        cfg['page_numbers'] = True
    if args.table_align:
        cfg['table_align'] = args.table_align
    if args.discipline:
        cfg['discipline'] = args.discipline
    if args.teacher:
        t_title, t_name = parse_teacher(args.teacher)
        cfg['teacher_title'] = t_title
        cfg['teacher_name'] = t_name
    if args.student:
        cfg['student_name'] = args.student

    if args.file:
        # Single file — metadata from CLI flags or config files lookup
        filename = args.file
        metadata = None
        if args.lab_type:
            metadata = (args.lab_type, args.num or '', args.topic or '')
        else:
            # Look up in config files list (normalize paths for comparison)
            norm_filename = os.path.normpath(filename)
            for entry in cfg.get('files', []):
                if os.path.normpath(entry['file']) == norm_filename:
                    metadata = (entry.get('type', ''), entry.get('num', ''), entry.get('topic', ''))
                    break
        print(f"  {filename}")
        convert_file(filename, metadata, src_dir, cfg=cfg)

    elif args.list:
        # External list file
        entries = load_list_file(args.list)
        for filename, meta in entries.items():
            print(f"  {filename}")
            convert_file(filename, meta, src_dir, cfg=cfg)

    else:
        # Default: convert all from config "files" list
        files = cfg.get('files', [])
        if not files:
            print("No files to convert. Add \"files\" to config or use -f/-l.")
            return
        for entry in files:
            filename = entry['file']
            metadata = (entry.get('type', ''), entry.get('num', ''), entry.get('topic', ''))
            print(f"  {filename}")
            convert_file(filename, metadata, src_dir, cfg=cfg)

    print("\nDone!")


if __name__ == "__main__":
    main()
